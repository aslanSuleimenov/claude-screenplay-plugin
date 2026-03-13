"""
Конвертер analytics/pitch.md → versions/pitch_vNN.docx

Запуск:
    python converter_MD_DOCX/pitch_to_docx.py
    python converter_MD_DOCX/pitch_to_docx.py МойПроект_pitch_v2

Зависимость: python-docx
    pip install python-docx
"""

import os
import re
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── Константы оформления ──────────────────────────────────────────────────────

FONT_MAIN   = "Georgia"
FONT_TITLE  = "Georgia"

PT_TITLE    = 24
PT_SUBTITLE = 11
PT_SECTION  = 13
PT_BODY     = 11
PT_LOGLINE  = 12
PT_META     = 10

COLOR_SECTION = RGBColor(0x1a, 0x1a, 0x2e)   # тёмно-синий
COLOR_META    = RGBColor(0x66, 0x66, 0x66)   # серый
COLOR_RULE    = RGBColor(0xcc, 0xcc, 0xcc)   # светло-серый


# ── Вспомогательные функции ───────────────────────────────────────────────────

def add_horizontal_rule(doc):
    """Тонкая горизонтальная линия."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "CCCCCC")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def set_run_font(run, size_pt, bold=False, italic=False, color=None):
    run.font.name  = FONT_MAIN
    run.font.size  = Pt(size_pt)
    run.font.bold  = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color


def add_para(doc, text, size_pt, bold=False, italic=False,
             align=WD_ALIGN_PARAGRAPH.LEFT, color=None,
             space_before=0, space_after=6):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    set_run_font(run, size_pt, bold, italic, color)
    return p


def parse_inline(paragraph, text, size_pt, base_italic=False):
    """Разбирает **bold** и *italic* внутри строки."""
    tokens = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*)', text)
    for token in tokens:
        if token.startswith("**") and token.endswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size_pt, bold=True, italic=base_italic)
        elif token.startswith("*") and token.endswith("*"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, size_pt, italic=True)
        elif token:
            run = paragraph.add_run(token)
            set_run_font(run, size_pt, italic=base_italic)


# ── Парсер pitch.md ───────────────────────────────────────────────────────────

def convert_pitch(source_path, output_path):
    text = source_path.read_text(encoding="utf-8")
    lines = text.splitlines()

    doc = Document()

    # Поля страницы
    for section in doc.sections:
        section.page_width   = Cm(21)
        section.page_height  = Cm(29.7)
        section.top_margin   = Cm(2.5)
        section.bottom_margin = Cm(2)
        section.left_margin  = Cm(3)
        section.right_margin = Cm(2.5)

    # Удалить дефолтный пустой абзац
    for el in list(doc.element.body):
        doc.element.body.remove(el)

    i = 0
    in_logline = False

    while i < len(lines):
        line = lines[i].rstrip()

        # ── Пустая строка ─────────────────────────────────────────────────────
        if line == "":
            i += 1
            continue

        # ── H1: Питч: Название ────────────────────────────────────────────────
        if line.startswith("# "):
            title = line[2:].strip()
            # убрать "Питч: " для выноса в подзаголовок
            if title.lower().startswith("питч:"):
                label  = "ПИТЧ"
                name   = title[5:].strip()
            else:
                label  = "ПИТЧ"
                name   = title

            doc.add_paragraph()  # отступ сверху
            add_para(doc, label, PT_META, bold=True, italic=False,
                     align=WD_ALIGN_PARAGRAPH.CENTER, color=COLOR_META,
                     space_before=0, space_after=2)
            add_para(doc, name, PT_TITLE, bold=True,
                     align=WD_ALIGN_PARAGRAPH.CENTER,
                     space_before=0, space_after=4)
            add_horizontal_rule(doc)
            i += 1
            continue

        # ── H2: Раздел ────────────────────────────────────────────────────────
        if line.startswith("## "):
            section_title = line[3:].strip().upper()
            in_logline = (section_title == "ЛОГЛАЙН")
            add_para(doc, section_title, PT_SECTION, bold=True,
                     color=COLOR_SECTION,
                     space_before=14, space_after=4)
            i += 1
            continue

        # ── Горизонтальная линия ---────────────────────────────────────────────
        if line.startswith("---"):
            add_horizontal_rule(doc)
            i += 1
            continue

        # ── Метаданные (жирная метка: значение) ───────────────────────────────
        meta_match = re.match(r'^\*\*([^*]+):\*\*\s*(.*)', line)
        if meta_match and not in_logline:
            label, value = meta_match.group(1), meta_match.group(2)
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(3)
            r1 = p.add_run(label + ": ")
            set_run_font(r1, PT_META, bold=True, color=COLOR_META)
            r2 = p.add_run(value)
            set_run_font(r2, PT_META, color=COLOR_META)
            i += 1
            continue

        # ── Логлайн (выделенный блок) ─────────────────────────────────────────
        if in_logline:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(6)
            p.paragraph_format.left_indent  = Cm(0.5)
            parse_inline(p, line, PT_LOGLINE, base_italic=True)
            i += 1
            continue

        # ── Обычный абзац ─────────────────────────────────────────────────────
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(5)
        parse_inline(p, line, PT_BODY)
        i += 1

    doc.save(str(output_path))
    print(f"Сохранено: {output_path}")


# ── main ──────────────────────────────────────────────────────────────────────

def main():
    root = Path(__file__).parent.parent
    source = root / "analytics" / "pitch.md"

    if not source.exists():
        print(f"Ошибка: файл {source} не найден.")
        print("Сначала запусти агент /pitch для генерации pitch.md")
        sys.exit(1)

    versions_dir = root / "versions"
    versions_dir.mkdir(exist_ok=True)

    # Определить имя файла
    if len(sys.argv) > 1:
        stem = sys.argv[1]
    else:
        project_name = root.name
        # Найти следующий номер версии
        existing = list(versions_dir.glob(f"{project_name}_pitch_v*.docx"))
        nums = []
        for f in existing:
            m = re.search(r'_pitch_v(\d+)\.docx$', f.name)
            if m:
                nums.append(int(m.group(1)))
        next_v = (max(nums) + 1) if nums else 1
        stem = f"{project_name}_pitch_v{next_v:02d}"

    output = versions_dir / f"{stem}.docx"
    convert_pitch(source, output)


if __name__ == "__main__":
    main()
